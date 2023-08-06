# -*- coding: utf-8 -*-
"""
@File  : table.py
@IDE   : PyCharm
@Author: Sanbom
@Date  : 2021/7/15
@Desc  : 
"""

from docx.document import Document as DocObject
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Cm
from docx.table import Table as DocTable, _Cell

from smartDocx.config import DocStylesBase
from smartDocx.constants import LENGTH_SWIFT, TableVerticalAlignment, TableLandscapeAlignment
from smartDocx.package.base import Base
from smartDocx.package.content import Content


class Table(Base):
    """表格"""

    def __init__(self, config: DocStylesBase, content_type: str = 'TABLE', *args, **kwargs):
        super(Table, self).__init__(content_type=content_type, config=config, *args, **kwargs)
        self.body_styles = self.styles
        if kwargs.get('head_config'):
            head_config = kwargs['head_config']
            self.head_styles = head_config if isinstance(head_config, dict) else head_config._to_dict()
        else:
            self.head_styles = self.body_styles
        self.column_width_dict = None
        self.table_head_width_dict = None
        self.table = None
        self.columns = None
        self.data = None
        self.max_merge_col_index = None
        self.doc_obj = None

    def generate(self, doc_obj: DocObject, columns: list, data: list, table_title: str = None,
                 table_title_style: DocStylesBase = None, built_table_color_style: str = '',
                 max_merge_col_index: int = None, horizontal_merge=True, vertical_merge=True):
        """
        生成表格
        :param doc_obj:
        :param columns: 表头列表
        :param data: 标题数据
        :param table_title: 表格标题
        :param table_title_style: 表格标题样式
        :param built_table_color_style: 内置表格样式
        :param max_merge_col_index: 合并单元格作用最大
        :param horizontal_merge: 水平方向合并单元格
        :param vertical_merge: 垂直方向合并单元格
        :return:
        """
        self.doc_obj = doc_obj
        self.columns = columns
        self.data = data
        self.max_merge_col_index = max_merge_col_index
        self.table_head_width_dict = {str(index): self._calculate_character_length(value) for index, value in
                                      enumerate(columns)}
        # 1.生成表标题
        if table_title:
            content = Content(config=table_title_style)
            content.generate(doc_obj, text=table_title)
        # 2.生成表格主体
        col_len = len(columns)
        row_len = len(data) + 1
        self.table = doc_obj.add_table(row_len, col_len)  # type: DocTable
        if built_table_color_style:
            self.table.style = built_table_color_style
        # 3.数据填充
        row_index = 0
        self.column_width_dict = {str(index): 0 for index in range(len(columns))}
        row_index = self._create_table_head(columns, row_index)
        self._fill_cell_data(data, row_index, horizontal_merge, vertical_merge)
        # 4.表格样式格式化
        self._format_table()
        self._justify_column_width()

    def _create_table_head(self, columns: list, row_index: int):
        """生成表头"""
        self.head_styles['font_weight'] = True
        self._fill_row_data(row=columns, row_index=row_index, is_table_head=True)
        return row_index + 1

    def _fill_cell_data(self, data: list, row_index: int, horizontal_merge: bool, vertical_merge: bool):
        """填充数据"""
        self.body_styles['font_weight'] = False
        for row in data:
            self._fill_row_data(row=row, row_index=row_index, horizontal_merge=horizontal_merge)
            # 垂直方向合并单元格
            if vertical_merge and self.max_merge_col_index is not None:
                self._merge_vertical_cells(row_index=row_index)
            row_index += 1

    def _fill_row_data(self, row: list, row_index: int, is_table_head: bool = False, horizontal_merge: bool = True):
        """填充行数据"""
        for col_index, value in enumerate(row):
            cell = self.table.cell(row_index, col_index)
            # 1.单元格样式设置
            self._format_cell_border(cell)
            run_obj = cell.paragraphs[0].add_run(str(value))
            # 2.单元格文字设置
            self._format_font(cell, run_obj, is_table_head)
            # 3.记录字符串长度
            value_len = self._calculate_character_length(str(value))
            if value_len > self.column_width_dict[str(col_index)]:
                self.column_width_dict[str(col_index)] = value_len
            # 4.水平方向合并单元格
            if horizontal_merge and self.max_merge_col_index is not None:
                self._merge_horizontal_cells(row_index=row_index, col_index=col_index)

        # 5.设置行高
        height = self.styles['font_size'] * LENGTH_SWIFT['pt'] * 2
        self.table.rows[row_index].height = Cm(height)

    def _merge_vertical_cells(self, row_index: int):
        """纵向合并单元格"""
        data_row_index = row_index - 1
        if row_index > 1:
            for col_index in range(0, self.max_merge_col_index + 1):
                if col_index > 0:
                    if self.data[data_row_index - 1][col_index - 1] != self.data[data_row_index][col_index - 1]:
                        continue
                pre = self.data[data_row_index - 1][col_index]
                cur = self.data[data_row_index][col_index]
                if cur and (pre == cur):
                    cell_ = self.table.cell(row_index - 1, col_index)  # 垂直方向上一个单元格
                    cell_text_ = cell_.paragraphs[0].text
                    cell = self.table.cell(row_index, col_index)
                    cell_text = cell.paragraphs[0].text
                    cell.paragraphs[0].text = ''  # 删除被合并单元格数据
                    try:
                        cell_.merge(cell)
                    except Exception as e:
                        # 合并单元格失败,恢复原始文本
                        cell_.paragraphs[0].text = cell_text_
                        cell.paragraphs[0].text = cell_text
                    else:
                        cell_.text = cell_.text.strip()

    def _merge_horizontal_cells(self, row_index: int, col_index: int):
        """横向合并单元格"""
        data_row_index = row_index - 1
        if 0 <= col_index <= self.max_merge_col_index:
            pre = self.data[data_row_index][col_index - 1]
            cur = self.data[data_row_index][col_index]
            if cur and (pre == cur):
                cell_ = self.table.cell(row_index, col_index - 1)  # 水平方向前一个单元格
                cell_text_ = cell_.paragraphs[0].text
                cell = self.table.cell(row_index, col_index)
                cell_text = cell.paragraphs[0].text
                cell.paragraphs[0].text = ''  # 删除被合并单元格数据
                try:
                    cell_.merge(cell)
                except Exception as e:
                    # 合并单元格失败,恢复原始文本
                    cell_.paragraphs[0].text = cell_text_
                    cell.paragraphs[0].text = cell_text
                else:
                    cell_.text = cell_.text.strip()

    def _format_table(self):
        """表格对齐"""
        for cell in self.table._cells:
            cell.paragraphs[0].paragraph_format.alignment = self.styles['alignment'] if self.styles[
                'alignment'] else TableLandscapeAlignment.center
            cell.paragraphs[0].paragraph_format.space_before = 0  # 段前设置
            cell.paragraphs[0].paragraph_format.space_after = 0  # 段后设置
            cell.paragraphs[0].paragraph_format.first_line_indent = Cm(0)  # 首航缩进为0
            cell.vertical_alignment = self.styles['vertical_alignment'] if self.styles[
                'vertical_alignment'] else TableVerticalAlignment.center
        self.table.alignment = self.styles['alignment']

    def _format_font(self, cell, run_obj, is_table_head: bool = False):
        """格式化文字"""
        font_styles = [
            'font_size',
            'font_weight',
            'color',
            'font_type',
            'italic',
            'underline',
            'background'
        ]
        if is_table_head:
            styles = self.head_styles  # 表头样式
        else:
            styles = self.body_styles  # 表体样式
        for font_style in font_styles:
            value = styles.get(font_style, None)
            handle_func = getattr(self, '_handle_{}'.format(font_style))
            if handle_func:
                if font_style == 'background':
                    handle_func(cell, value)
                else:
                    handle_func(run_obj, value)

    def _handle_background(self, cell, value):
        """单元格背景色设置"""
        if value:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), value))
            cell._tc.get_or_add_tcPr().append(shading_elm_1)

    @staticmethod
    def _format_cell_border(cell: _Cell, **kwargs):
        """
        设置单元格边框函数
        使用方法:
        Set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        传入参数有cell, 即单元格；top指上边框；bottom指下边框；start指左边框；end指右边框。
        "sz"指线的粗细程度；"val"指线型，比如单线，虚线等；"color"指颜色，颜色编码可百度；
        "space"指间隔，一般不设置，设置的值大于0会导致线错开；"shadow"指边框阴影
        """
        default_style = {"sz": 5, "color": "#000000", "val": "single"}
        if not kwargs.get('top'):
            kwargs['top'] = default_style
        if not kwargs.get('bottom'):
            kwargs['bottom'] = default_style
        if not kwargs.get('start'):
            kwargs['start'] = default_style
        if not kwargs.get('end'):
            kwargs['end'] = default_style

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def _justify_column_width(self):
        """表格宽度自适应 会重新设置字体样式"""
        self.table.autofit = False  # 关闭自带的自适应
        self._calculate_appropriate_column_width()  # 计算合适的列宽度
        for row_index in range(len(self.table.rows)):
            for col_index in range(len(self.columns)):
                cell = self.table.cell(row_index, col_index)
                width = self.column_width_dict[str(col_index)]
                cell.width = Cm(width)
                if row_index == 0:
                    self._format_font(cell=cell, run_obj=cell.paragraphs[0].runs[0], is_table_head=True)
                else:
                    self._format_font(cell=cell, run_obj=cell.paragraphs[0].runs[0])

    def _calculate_character_length(self, character):
        """
        计算字符串长度
        :param character:
        :return:
        """
        character_len = len(str(character))
        unicode_len = len(str(character).encode('utf-8'))
        if character_len == unicode_len:
            # ASCII码字符组成
            # Word文本字符长度
            appearance_len = round(character_len * self.styles['font_size'] * LENGTH_SWIFT['pt'] / 2, 2)
        else:
            # 非ASCII和ASCII字符混合
            # 汉字Unicode编码占3个字节,ASCII码占1个字节
            ascii_len = (unicode_len - character_len) / 2
            noascii_len = (unicode_len - ascii_len) / 3
            # Word文本字符长度
            appearance_len = round(
                character_len * self.styles['font_size'] * LENGTH_SWIFT['pt'] / 2 + noascii_len * self.styles[
                    'font_size'] * LENGTH_SWIFT['pt'], 2)
        return round(appearance_len * 1.3, 2)  # 考虑到表格单元格内边距

    def _calculate_appropriate_column_width(self):
        """计算合适的单元格宽度"""
        page_width_cm = self.doc_obj.sections[-1].page_width.cm
        page_left_cm = self.doc_obj.sections[-1].left_margin.cm
        page_right_cm = self.doc_obj.sections[-1].right_margin.cm
        TABLE_TOTAL_WIDTH = page_width_cm - page_right_cm - page_left_cm  # 表格宽度
        col_width_sum = sum(self.column_width_dict.values())
        col_num = len(self.column_width_dict)
        col_width_avg = round(col_width_sum / col_num, 3)
        redundant_len = TABLE_TOTAL_WIDTH - col_width_sum
        if redundant_len < 0:
            col_great_then_avg = {k: v for k, v in self.column_width_dict.items() if v > col_width_avg}
            col_less_and_equal_then_avg = {k: v for k, v in self.column_width_dict.items() if
                                           v <= col_width_avg}
            sum_great_then_avg = sum(col_great_then_avg.values())
            sum_less_and_equal_then_avg = sum(col_less_and_equal_then_avg.values())
            redundant_len_ = TABLE_TOTAL_WIDTH - sum_less_and_equal_then_avg
            if redundant_len_ >= 0:
                for k, v in col_great_then_avg.items():
                    self.column_width_dict[k] = (v / sum_great_then_avg * redundant_len_)
            else:
                for k, v in self.column_width_dict.items():
                    self.column_width_dict[k] = (v / col_width_sum * TABLE_TOTAL_WIDTH)
        else:
            for k, v in self.column_width_dict.items():
                self.column_width_dict[k] = (v / col_width_sum * redundant_len) + v
