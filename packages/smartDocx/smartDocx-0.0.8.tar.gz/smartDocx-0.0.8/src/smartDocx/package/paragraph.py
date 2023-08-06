# -*- coding: utf-8 -*-
"""
@File  : paragraph.py
@IDE   : PyCharm
@Author: Sanbom
@Date  : 2021/7/13
@Desc  : 
"""

import re

from docx import Document
from docx.document import Document as DocObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from smartDocx.config import DocStylesBase
from smartDocx.package.base import Base


class Paragraph(Base):
    """段落"""

    def __init__(self, config: DocStylesBase, content_type: str = 'PARAGRAPH', *args, **kwargs):
        super(Paragraph, self).__init__(config=config, content_type=content_type, *args, **kwargs)
        self._run_obj = None

    def generate(self, doc_obj: DocObject, text: str, append=False):
        """生成段落"""
        if append:
            # 去除上一个run_obj的换行符号
            paragraph = doc_obj.paragraphs[-1]
            prev_run_obj = paragraph.runs[-1]
            prev_run_obj.text = re.sub('\t', '', prev_run_obj.text)
            prev_run_obj.text = re.sub('\n', '', prev_run_obj.text)
        else:
            paragraph = doc_obj.add_paragraph()  # 添加新段落
        self._run_obj = paragraph.add_run(text=text)  # 新段落添加文本
        self._format(paragraph, self._run_obj)  # 字体格式化

    def _format(self, paragraph, run_obj):
        """格式化"""
        # 新段落格式化
        self._format_paragraph(paragraph)
        # 字体格式化
        self._format_font(run_obj=run_obj)

    def _format_paragraph(self, paragraph):
        """段落样式设置"""
        paragraph_styles = [
            'text_align',
            'line_indent',
            'line_spacing',
            'space_before',
            'space_after',
            'paragraph_underline'
        ]
        for paragraph_style in paragraph_styles:
            value = self.styles.get(paragraph_style, None)
            handle_func = getattr(self, '_handle_{}'.format(paragraph_style))
            if handle_func:
                handle_func(paragraph, value)

    def _format_font(self, run_obj):
        """字体样式设置"""
        font_styles = [
            'font_size',
            'font_weight',
            'color',
            'font_type',
            'italic',
            'underline'
        ]
        for font_style in font_styles:
            value = self.styles.get(font_style, None)
            handle_func = getattr(self, '_handle_{}'.format(font_style))
            if handle_func:
                handle_func(run_obj, value)

    # ==============
    # = 段落格式设置 =
    # ==============
    def _handle_line_indent(self, obj, value):
        """设置首航缩进"""
        if value is True:
            obj.paragraph_format.first_line_indent = Inches(0.4)
        else:
            obj.paragraph_format.first_line_indent = Inches(0)

    def _handle_line_spacing(self, obj, value):
        """设置行间距"""
        if isinstance(value, (int, float)):
            obj.paragraph_format.line_spacing = value

    def _handle_space_before(self, obj, value):
        """设置段前"""
        if isinstance(value, (int, float)):
            obj.paragraph_format.space_before = Pt(value)

    def _handle_space_after(self, obj, value):
        """设置段后"""
        if isinstance(value, (int, float)):
            obj.paragraph_format.space_after = Pt(value)

    def _handle_text_align(self, obj, value):
        """设置对齐方式"""
        obj.alignment = value

    def _handle_paragraph_underline(self, paragraph, value):
        """段落底部下划线设置"""
        if value is None:
            return

        prPr = paragraph._element.get_or_add_pPr()
        kwargs = dict()
        kwargs['bottom'] = value

        tcBorders = prPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            prPr.append(tcBorders)

        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))
