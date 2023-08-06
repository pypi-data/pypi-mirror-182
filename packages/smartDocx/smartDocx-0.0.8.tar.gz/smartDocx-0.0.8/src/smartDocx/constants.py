"""
@File  : constants.py
@IDE   : PyCharm
@Author: Sanbom
@Date  : 2021/7/23
@Desc  : 
"""
# -*- coding: utf-8 -*-
import os
from enum import Enum

from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 根目录
BASE_DIR = os.path.dirname(os.path.abspath(__file__)).replace('\\', '/')
TMP_DIR = os.path.join(BASE_DIR, 'tmp').replace('\\', '/')
TEMPLATE_DIR = os.path.join(BASE_DIR, 'template').replace('\\', '/')
DATA_DIR = os.path.join(BASE_DIR, 'data').replace('\\', '/')


class DocFontSize(Enum):
    """文档字体磅值"""
    Zero = 36  # 初号
    SmallZero = 31  # 小初号
    One = 27  # 一号
    SmallOne = 24  # 小一号
    Two = 21
    SmallTwo = 18
    Three = 16
    SmallThree = 15
    Four = 14
    SmallFour = 12
    Five = 10.5
    SmallFive = 9
    Six = 8
    SmallSix = 7
    Seven = 6
    SmallSeven = 5


# 长度单位转换
LENGTH_SWIFT = {
    'inches': 2.54,  # 英寸
    'cm': 1,  # 厘米
    'mm': 0.1,  # 毫米
    'pt': 0.0376,  # 点或磅
    'twips': 1 / 567000,  # 缇
}


class DocFontFamily(Enum):
    """中文字体格式"""
    ST = '宋体'
    HT = '黑体'
    MSYH = '微软雅黑'
    KT = '楷体'


class Orientation(Enum):
    """页面排布方向"""
    # 纸张方向: 默认纵向(PORTRAIT), 横向(LANDSCAPE)
    landscape = WD_ORIENTATION.LANDSCAPE
    portrait = WD_ORIENTATION.PORTRAIT


class DocAlignment(Enum):
    """文本水平对齐方式"""
    left = WD_PARAGRAPH_ALIGNMENT.LEFT
    right = WD_PARAGRAPH_ALIGNMENT.RIGHT
    center = WD_PARAGRAPH_ALIGNMENT.CENTER
    justify = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


class TableLandscapeAlignment(Enum):
    """表格单元格文本, 水平对齐方式"""
    left = WD_TABLE_ALIGNMENT.LEFT
    right = WD_TABLE_ALIGNMENT.RIGHT
    center = WD_TABLE_ALIGNMENT.CENTER


class TableVerticalAlignment(Enum):
    """表格单元格文本,垂直对齐方式"""
    top = WD_CELL_VERTICAL_ALIGNMENT.TOP
    bottom = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    center = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# 默认表格样式
class WordBuiltinStyle(object):
    """
    word默认表格背景色样式
        样式对应效果,请查看template文件夹下的table_styles.docx文件

        #### 生成代码如下
        import docx

        from smart_smart_report.config import WordTableBackgroundStyle
        document = docx.Document()
        for style_ in document.styles:
            name = style_.name
            if name in WordTableBackgroundStyle()._data:
                document.add_paragraph(text=name)
                table = document.add_table(7, 3, style=style_)
                document.add_paragraph()
        document.save('xxx.docx')

    """

    def __init__(self):
        self._data = {
            "Normal": {'type_name': 'paragraph', 'name': 'Normal', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Header": {'type_name': 'paragraph', 'name': 'Header', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Footer": {'type_name': 'paragraph', 'name': 'Footer', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Heading 1": {'type_name': 'paragraph', 'name': 'Heading 1', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 2": {'type_name': 'paragraph', 'name': 'Heading 2', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 3": {'type_name': 'paragraph', 'name': 'Heading 3', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 4": {'type_name': 'paragraph', 'name': 'Heading 4', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 5": {'type_name': 'paragraph', 'name': 'Heading 5', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 6": {'type_name': 'paragraph', 'name': 'Heading 6', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 7": {'type_name': 'paragraph', 'name': 'Heading 7', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 8": {'type_name': 'paragraph', 'name': 'Heading 8', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Heading 9": {'type_name': 'paragraph', 'name': 'Heading 9', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "No Spacing": {'type_name': 'paragraph', 'name': 'No Spacing', 'type': WD_STYLE_TYPE.PARAGRAPH,
                           'zh_name': ''},
            "Title": {'type_name': 'paragraph', 'name': 'Title', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Subtitle": {'type_name': 'paragraph', 'name': 'Subtitle', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "List Paragraph": {'type_name': 'paragraph', 'name': 'List Paragraph', 'type': WD_STYLE_TYPE.PARAGRAPH,
                               'zh_name': ''},
            "Body Text": {'type_name': 'paragraph', 'name': 'Body Text', 'type': WD_STYLE_TYPE.PARAGRAPH,
                          'zh_name': ''},
            "Body Text 2": {'type_name': 'paragraph', 'name': 'Body Text 2', 'type': WD_STYLE_TYPE.PARAGRAPH,
                            'zh_name': ''},
            "Body Text 3": {'type_name': 'paragraph', 'name': 'Body Text 3', 'type': WD_STYLE_TYPE.PARAGRAPH,
                            'zh_name': ''},
            "List": {'type_name': 'paragraph', 'name': 'List', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "List 2": {'type_name': 'paragraph', 'name': 'List 2', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "List 3": {'type_name': 'paragraph', 'name': 'List 3', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "List Bullet": {'type_name': 'paragraph', 'name': 'List Bullet', 'type': WD_STYLE_TYPE.PARAGRAPH,
                            'zh_name': ''},
            "List Bullet 2": {'type_name': 'paragraph', 'name': 'List Bullet 2', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "List Bullet 3": {'type_name': 'paragraph', 'name': 'List Bullet 3', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "List Number": {'type_name': 'paragraph', 'name': 'List Number', 'type': WD_STYLE_TYPE.PARAGRAPH,
                            'zh_name': ''},
            "List Number 2": {'type_name': 'paragraph', 'name': 'List Number 2', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "List Number 3": {'type_name': 'paragraph', 'name': 'List Number 3', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "List Continue": {'type_name': 'paragraph', 'name': 'List Continue', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "List Continue 2": {'type_name': 'paragraph', 'name': 'List Continue 2', 'type': WD_STYLE_TYPE.PARAGRAPH,
                                'zh_name': ''},
            "List Continue 3": {'type_name': 'paragraph', 'name': 'List Continue 3', 'type': WD_STYLE_TYPE.PARAGRAPH,
                                'zh_name': ''},
            "macro": {'type_name': 'paragraph', 'name': 'macro', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Quote": {'type_name': 'paragraph', 'name': 'Quote', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Caption": {'type_name': 'paragraph', 'name': 'Caption', 'type': WD_STYLE_TYPE.PARAGRAPH, 'zh_name': ''},
            "Intense Quote": {'type_name': 'paragraph', 'name': 'Intense Quote', 'type': WD_STYLE_TYPE.PARAGRAPH,
                              'zh_name': ''},
            "TOC Heading": {'type_name': 'paragraph', 'name': 'TOC Heading', 'type': WD_STYLE_TYPE.PARAGRAPH,
                            'zh_name': ''},
            "Header Char": {'type_name': 'character', 'name': 'Header Char', 'type': WD_STYLE_TYPE.CHARACTER,
                            'zh_name': ''},
            "Footer Char": {'type_name': 'character', 'name': 'Footer Char', 'type': WD_STYLE_TYPE.CHARACTER,
                            'zh_name': ''},
            "Default Paragraph Font": {'type_name': 'character', 'name': 'Default Paragraph Font',
                                       'CHARACTER': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Heading 1 Char": {'type_name': 'character', 'name': 'Heading 1 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 2 Char": {'type_name': 'character', 'name': 'Heading 2 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 3 Char": {'type_name': 'character', 'name': 'Heading 3 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Title Char": {'type_name': 'character', 'name': 'Title Char', 'type': WD_STYLE_TYPE.CHARACTER,
                           'zh_name': ''},
            "Subtitle Char": {'type_name': 'character', 'name': 'Subtitle Char', 'type': WD_STYLE_TYPE.CHARACTER,
                              'zh_name': ''},
            "Body Text Char": {'type_name': 'character', 'name': 'Body Text Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Body Text 2 Char": {'type_name': 'character', 'name': 'Body Text 2 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                                 'zh_name': ''},
            "Body Text 3 Char": {'type_name': 'character', 'name': 'Body Text 3 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                                 'zh_name': ''},
            "Macro Text Char": {'type_name': 'character', 'name': 'Macro Text Char', 'type': WD_STYLE_TYPE.CHARACTER,
                                'zh_name': ''},
            "Quote Char": {'type_name': 'character', 'name': 'Quote Char', 'type': WD_STYLE_TYPE.CHARACTER,
                           'zh_name': ''},
            "Heading 4 Char": {'type_name': 'character', 'name': 'Heading 4 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 5 Char": {'type_name': 'character', 'name': 'Heading 5 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 6 Char": {'type_name': 'character', 'name': 'Heading 6 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 7 Char": {'type_name': 'character', 'name': 'Heading 7 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 8 Char": {'type_name': 'character', 'name': 'Heading 8 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Heading 9 Char": {'type_name': 'character', 'name': 'Heading 9 Char', 'type': WD_STYLE_TYPE.CHARACTER,
                               'zh_name': ''},
            "Strong": {'type_name': 'character', 'name': 'Strong', 'type': WD_STYLE_TYPE.CHARACTER, 'zh_name': ''},
            "Emphasis": {'type_name': 'character', 'name': 'Emphasis', 'type': WD_STYLE_TYPE.CHARACTER, 'zh_name': ''},
            "Intense Quote Char": {'type_name': 'character', 'name': 'Intense Quote Char',
                                   'type': WD_STYLE_TYPE.CHARACTER, 'zh_name': ''},
            "Subtle Emphasis": {'type_name': 'character', 'name': 'Subtle Emphasis', 'type': WD_STYLE_TYPE.CHARACTER,
                                'zh_name': ''},
            "Intense Emphasis": {'type_name': 'character', 'name': 'Intense Emphasis', 'type': WD_STYLE_TYPE.CHARACTER,
                                 'zh_name': ''},
            "Subtle Reference": {'type_name': 'character', 'name': 'Subtle Reference', 'type': WD_STYLE_TYPE.CHARACTER,
                                 'zh_name': ''},
            "Intense Reference": {'type_name': 'character', 'name': 'Intense Reference',
                                  'type': WD_STYLE_TYPE.CHARACTER, 'zh_name': ''},
            "Book Title": {'type_name': 'character', 'name': 'Book Title', 'type': WD_STYLE_TYPE.CHARACTER,
                           'zh_name': ''},
            "Normal Table": {'type_name': 'table', 'name': 'Normal Table', 'type': WD_STYLE_TYPE.TABLE, 'zh_name': ''},
            "Table Grid": {'type_name': 'table', 'name': 'Table Grid', 'type': WD_STYLE_TYPE.TABLE, 'zh_name': ''},
            "Light Shading": {'type_name': 'table', 'name': 'Light Shading', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Light Shading Accent 1": {'type_name': 'table', 'name': 'Light Shading Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light Shading Accent 2": {'type_name': 'table', 'name': 'Light Shading Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light Shading Accent 3": {'type_name': 'table', 'name': 'Light Shading Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light Shading Accent 4": {'type_name': 'table', 'name': 'Light Shading Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light Shading Accent 5": {'type_name': 'table', 'name': 'Light Shading Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light Shading Accent 6": {'type_name': 'table', 'name': 'Light Shading Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Light List": {'type_name': 'table', 'name': 'Light List', 'type': WD_STYLE_TYPE.TABLE, 'zh_name': ''},
            "Light List Accent 1": {'type_name': 'table', 'name': 'Light List Accent 1', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light List Accent 2": {'type_name': 'table', 'name': 'Light List Accent 2', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light List Accent 3": {'type_name': 'table', 'name': 'Light List Accent 3', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light List Accent 4": {'type_name': 'table', 'name': 'Light List Accent 4', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light List Accent 5": {'type_name': 'table', 'name': 'Light List Accent 5', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light List Accent 6": {'type_name': 'table', 'name': 'Light List Accent 6', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid": {'type_name': 'table', 'name': 'Light Grid', 'type': WD_STYLE_TYPE.TABLE, 'zh_name': ''},
            "Light Grid Accent 1": {'type_name': 'table', 'name': 'Light Grid Accent 1', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid Accent 2": {'type_name': 'table', 'name': 'Light Grid Accent 2', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid Accent 3": {'type_name': 'table', 'name': 'Light Grid Accent 3', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid Accent 4": {'type_name': 'table', 'name': 'Light Grid Accent 4', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid Accent 5": {'type_name': 'table', 'name': 'Light Grid Accent 5', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Light Grid Accent 6": {'type_name': 'table', 'name': 'Light Grid Accent 6', 'type': WD_STYLE_TYPE.TABLE,
                                    'zh_name': ''},
            "Medium Shading 1": {'type_name': 'table', 'name': 'Medium Shading 1', 'type': WD_STYLE_TYPE.TABLE,
                                 'zh_name': ''},
            "Medium Shading 1 Accent 1": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 1',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 1 Accent 2": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 2',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 1 Accent 3": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 3',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 1 Accent 4": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 4',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 1 Accent 5": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 5',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 1 Accent 6": {'type_name': 'table', 'name': 'Medium Shading 1 Accent 6',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2": {'type_name': 'table', 'name': 'Medium Shading 2', 'type': WD_STYLE_TYPE.TABLE,
                                 'zh_name': ''},
            "Medium Shading 2 Accent 1": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 1',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2 Accent 2": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 2',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2 Accent 3": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 3',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2 Accent 4": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 4',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2 Accent 5": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 5',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium Shading 2 Accent 6": {'type_name': 'table', 'name': 'Medium Shading 2 Accent 6',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Medium List 1": {'type_name': 'table', 'name': 'Medium List 1', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Medium List 1 Accent 1": {'type_name': 'table', 'name': 'Medium List 1 Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 1 Accent 2": {'type_name': 'table', 'name': 'Medium List 1 Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 1 Accent 3": {'type_name': 'table', 'name': 'Medium List 1 Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 1 Accent 4": {'type_name': 'table', 'name': 'Medium List 1 Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 1 Accent 5": {'type_name': 'table', 'name': 'Medium List 1 Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 1 Accent 6": {'type_name': 'table', 'name': 'Medium List 1 Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2": {'type_name': 'table', 'name': 'Medium List 2', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Medium List 2 Accent 1": {'type_name': 'table', 'name': 'Medium List 2 Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2 Accent 2": {'type_name': 'table', 'name': 'Medium List 2 Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2 Accent 3": {'type_name': 'table', 'name': 'Medium List 2 Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2 Accent 4": {'type_name': 'table', 'name': 'Medium List 2 Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2 Accent 5": {'type_name': 'table', 'name': 'Medium List 2 Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium List 2 Accent 6": {'type_name': 'table', 'name': 'Medium List 2 Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1": {'type_name': 'table', 'name': 'Medium Grid 1', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Medium Grid 1 Accent 1": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1 Accent 2": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1 Accent 3": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1 Accent 4": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1 Accent 5": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 1 Accent 6": {'type_name': 'table', 'name': 'Medium Grid 1 Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2": {'type_name': 'table', 'name': 'Medium Grid 2', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Medium Grid 2 Accent 1": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2 Accent 2": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2 Accent 3": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2 Accent 4": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2 Accent 5": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 2 Accent 6": {'type_name': 'table', 'name': 'Medium Grid 2 Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3": {'type_name': 'table', 'name': 'Medium Grid 3', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Medium Grid 3 Accent 1": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3 Accent 2": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3 Accent 3": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3 Accent 4": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3 Accent 5": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Medium Grid 3 Accent 6": {'type_name': 'table', 'name': 'Medium Grid 3 Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Dark List": {'type_name': 'table', 'name': 'Dark List', 'type': WD_STYLE_TYPE.TABLE, 'zh_name': ''},
            "Dark List Accent 1": {'type_name': 'table', 'name': 'Dark List Accent 1', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Dark List Accent 2": {'type_name': 'table', 'name': 'Dark List Accent 2', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Dark List Accent 3": {'type_name': 'table', 'name': 'Dark List Accent 3', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Dark List Accent 4": {'type_name': 'table', 'name': 'Dark List Accent 4', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Dark List Accent 5": {'type_name': 'table', 'name': 'Dark List Accent 5', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Dark List Accent 6": {'type_name': 'table', 'name': 'Dark List Accent 6', 'type': WD_STYLE_TYPE.TABLE,
                                   'zh_name': ''},
            "Colorful Shading": {'type_name': 'table', 'name': 'Colorful Shading', 'type': WD_STYLE_TYPE.TABLE,
                                 'zh_name': ''},
            "Colorful Shading Accent 1": {'type_name': 'table', 'name': 'Colorful Shading Accent 1',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful Shading Accent 2": {'type_name': 'table', 'name': 'Colorful Shading Accent 2',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful Shading Accent 3": {'type_name': 'table', 'name': 'Colorful Shading Accent 3',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful Shading Accent 4": {'type_name': 'table', 'name': 'Colorful Shading Accent 4',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful Shading Accent 5": {'type_name': 'table', 'name': 'Colorful Shading Accent 5',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful Shading Accent 6": {'type_name': 'table', 'name': 'Colorful Shading Accent 6',
                                          'type': WD_STYLE_TYPE.TABLE,
                                          'zh_name': ''},
            "Colorful List": {'type_name': 'table', 'name': 'Colorful List', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Colorful List Accent 1": {'type_name': 'table', 'name': 'Colorful List Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful List Accent 2": {'type_name': 'table', 'name': 'Colorful List Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful List Accent 3": {'type_name': 'table', 'name': 'Colorful List Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful List Accent 4": {'type_name': 'table', 'name': 'Colorful List Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful List Accent 5": {'type_name': 'table', 'name': 'Colorful List Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful List Accent 6": {'type_name': 'table', 'name': 'Colorful List Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid": {'type_name': 'table', 'name': 'Colorful Grid', 'type': WD_STYLE_TYPE.TABLE,
                              'zh_name': ''},
            "Colorful Grid Accent 1": {'type_name': 'table', 'name': 'Colorful Grid Accent 1',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid Accent 2": {'type_name': 'table', 'name': 'Colorful Grid Accent 2',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid Accent 3": {'type_name': 'table', 'name': 'Colorful Grid Accent 3',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid Accent 4": {'type_name': 'table', 'name': 'Colorful Grid Accent 4',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid Accent 5": {'type_name': 'table', 'name': 'Colorful Grid Accent 5',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "Colorful Grid Accent 6": {'type_name': 'table', 'name': 'Colorful Grid Accent 6',
                                       'type': WD_STYLE_TYPE.TABLE,
                                       'zh_name': ''},
            "No List": {'type_name': 'numbering', 'name': 'No List', 'type': WD_STYLE_TYPE.LIST, 'zh_name': ''}
        }

    @property
    def data(self):
        return self._data

    def __contains__(self, item):
        if item in self._data:
            return True
        return False

    def __getitem__(self, item):
        if item in self:
            return self._data[item]
        raise KeyError('无该类别样式: '.format(item))

    def __iter__(self):
        for item in self.data:
            yield item


WORD_BUILTIN_STYLES = WordBuiltinStyle()

# 默认英文字体
DEFAULT_WEST_FONT = 'Times New Roman'

if __name__ == '__main__':
    center = DocAlignment.center
    print(center.value)
