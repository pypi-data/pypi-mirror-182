"""
@File  : report.py
@IDE   : PyCharm
@Author: Sanbom
@Date  : 2021/7/13
@Desc  : 
"""
# -*- coding: utf-8 -*-
import datetime
import os
import re
import sys

from docx import Document
from docx.document import Document as DocObject
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, ns
from docx.shared import Cm

from smartDocx.config import PaperTitleStyleBase, PaperSubTitleStyleBase, AbstractStyleBase, \
    FirstTitleStyleBase, SecondTitleStyleBase, ThirdTitleStyleBase, FourthTitleStyleBase, ContentStyleBase, \
    TableTitleStyleBase, TableHeadStyleBase, TableBodyStyleBase, PicTitleStyleBase, PageStyleBase
from smartDocx.constants import WORD_BUILTIN_STYLES, TEMPLATE_DIR, Orientation


class SmartReportDataBase(object):
    """智能报告数据"""

    def __init__(self, sql_client=None, neo4j_client=None, es_client=None, *args, **kwargs):
        self._data = {
            'global': dict(
                sql_client=sql_client,
                neo4j_client=neo4j_client,
                es_client=es_client
            )
        }
        self.sql_client = sql_client
        self.neo4j_client = neo4j_client
        self.es_client = es_client
        self.args = args
        self.kwargs = kwargs

    @property
    def data(self):
        """获取报告数据"""
        self.distribute()
        return self._data

    def distribute(self):
        """派发任务"""
        for attr in dir(self):
            if attr.startswith('chapter_'):
                handler = getattr(self, attr)
                handler()

    def get_chapter_level(self, chapter_function_name: str):
        """获取章节层级"""
        return chapter_function_name.split('chapter_')[1]


class SmartReportBase(object):
    """智能报告"""

    def __init__(self, filename: str, data: dict, template: str = None, *args, **kwargs):
        """
        智能报告

        :param filename: 文件绝对路径
        :param data: 数据
        :param args:
        :param kwargs:
        """
        self.args = args
        self.kwargs = kwargs
        if not os.path.isabs(filename):
            filename = os.path.abspath(filename)
        self.filename = filename
        self.data = data
        if template:
            # 读取模板文件, 路径检索顺序: 模板文件夹>全局检索
            if not os.path.exists(os.path.join(TEMPLATE_DIR, template)):
                if not os.path.exists(template):
                    raise FileNotFoundError(template)
            else:
                template = os.path.join(TEMPLATE_DIR, template)
            print('从模板文件创建: {}'.format(template))
            self.doc_obj = Document(template)  # type: DocObject
        else:
            # 创建新文件
            self.doc_obj = Document()  # type: DocObject
            print('从新文件创建: {}'.format(template))
        self._clear_files = set()
        # 页面布局(cm):默认A4纸大小
        self.page_conf = PageStyleBase()
        self.default_orientation = self.page_conf['orientation']
        self._load_builtin_style()
        self._handle_filename()

        # 默认样式设置
        # 文章标题
        self.title_level_0_style = PaperTitleStyleBase()
        # 文章副标题样式
        self.paper_subheading_style = PaperSubTitleStyleBase()
        # 文章概要
        self.abstract = AbstractStyleBase()
        # 一级标题:段落标题
        self.title_level_1_style = FirstTitleStyleBase()
        # 二级标题:
        self.title_level_2_style = SecondTitleStyleBase()
        # 三级标题
        self.title_level_3_style = ThirdTitleStyleBase()
        # 四级标题
        self.title_level_4_style = FourthTitleStyleBase()
        # 正文样式
        self.content_style = ContentStyleBase()
        # 表格样式
        self.table_title_style = TableTitleStyleBase()  # 标题
        self.table_head_style = TableHeadStyleBase()  # 表头
        self.table_body_style = TableBodyStyleBase()  # 表体
        # 图片样式
        self.picture_title_style = PicTitleStyleBase()

    def generate(self):
        """生成"""
        print('=' * 100)
        print('自动化报告开始: {}'.format(self.filename))
        # 文章段落结构排序
        self._sort_chapter()
        print('文章段落结排序结束!')
        # 页面设置
        self._handle_page_settings()
        print('页面设置初始化成功!')
        # 目录
        # python脚本无法直接实现, 模板文件事先准备好目录部分, 完成后打开文件直接刷新.
        # 文章标题
        self.handle_paper_heading()
        print('页面标题写入成功!')
        # 文章副标题
        self.handle_paper_subheading()
        print('页面副标题写入成功!')
        # 文章摘要
        self.handle_abstract()
        print('文章摘要写入成功!')
        # 文章段落内容
        self._handle_chapter()
        print('文章段落内容处理开始!')
        # 页眉和页脚
        # 模板文件事先准备好
        # 保存
        self.doc_obj.save(self.filename)
        print('报告保存本地成功!')
        # 删除临时文件
        self._clear()
        print('清理缓存文件结束!')
        print('输出文件: {}'.format(self.filename))

    # ===============
    # = 子类需重写方法 =
    # ===============
    def handle_paper_heading(self):
        """文章标题"""
        pass

    def handle_paper_subheading(self):
        """文章副标题"""
        pass

    def handle_abstract(self):
        """文章摘要"""
        pass

    # ===============
    # = 子类可调用方法 =
    # ===============
    @staticmethod
    def delete_paragraph(paragraph):
        """删除段落"""
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def rotate_page_orientation(self, orientation: Orientation = None):
        """旋转页面方向"""
        section = self.doc_obj.add_section()
        if orientation is None:
            # 页面方向缺失, 恢复默认方向
            orientation = self.default_orientation
        self._set_orientation(section, orientation)

    def lines_break(self, lines: int):
        """换行"""
        for line in range(lines):
            self.doc_obj.add_paragraph()

    def page_break(self):
        """换页"""
        paragraph = self.doc_obj.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def clear_file(self, file):
        """"""
        self._clear_files.add(file)

    # =============
    # == 样式设置 ==
    # =============
    def _load_builtin_style(self):
        """加载所有Word内置样式"""
        for style in WORD_BUILTIN_STYLES:
            item = WORD_BUILTIN_STYLES[style]
            try:
                self.doc_obj.styles.add_style(style, item['type'], builtin=True)
            except:
                pass
        # 设置正文首行缩进为0,防止读取模板情况下,新增内容设置无效!
        self.doc_obj.styles['Normal'].paragraph_format.first_line_indent = Cm(0)

    def _handle_page_settings(self):
        """页面设置"""
        for section in self.doc_obj.sections:
            self._set_orientation(section, self.default_orientation)
            section.left_margin = Cm(self.page_conf['left_margin'])
            section.right_margin = Cm(self.page_conf['right_margin'])
            section.top_margin = Cm(self.page_conf['top_margin'])
            section.bottom_margin = Cm(self.page_conf['bottom_margin'])
            section.page_width = Cm(self.page_conf['page_width'])
            section.page_height = Cm(self.page_conf['page_height'])

    def _set_orientation(self, section, orientation: Orientation = None):
        """页面旋转"""
        if isinstance(orientation, (Orientation,)):
            orientation = orientation.value
        section.orientation = orientation
        # 必须同时设置方向、新宽度和高度值,才会生效
        if orientation == WD_ORIENTATION.PORTRAIT:
            section.page_width = Cm(self.page_conf['page_width'])
            section.page_height = Cm(self.page_conf['page_height'])
        else:
            section.page_height = Cm(self.page_conf['page_width'])
            section.page_width = Cm(self.page_conf['page_height'])

    def _handle_header(self):
        """设置页眉"""
        # 模板自行设置
        pass

    def _handle_footer(self):
        """设置页脚"""
        # 模板自行设置
        # footer = self.doc_obj.sections[0].footer
        # self._add_page_number(footer.paragraphs[0].add_run())
        # footer.paragraphs[0].alignment = ALIGNMENT_DICT['center']
        pass

    def _add_page_number(self, run):
        """页脚添加页码"""
        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    @staticmethod
    def _create_element(name):
        """创建Oxml元素"""
        return OxmlElement(name)

    @staticmethod
    def _create_attribute(element, name, value):
        """添加页码属性"""
        element.set(ns.qn(name), value)

    # =====================
    # == 章节及段落数据处理 ==
    # =====================
    def _sort_chapter(self):
        """"""
        regex_rules = [
            r'(chapter_\d{2})_(paragraph_\d{2})',
            r'(chapter_\d{2}_\d{2})_(paragraph_\d{2})',
            r'(chapter_\d{2}_\d{2}_\d{2})_(paragraph_\d{2})',
            r'(chapter_\d{2}_\d{2}_\d{2}_\d{2})_(paragraph_\d{2})',
        ]
        chapter_list = list()
        subchapter_dict = dict()
        paragraph_dict = dict()
        for attr in dir(self):
            if attr.startswith('chapter'):
                # 处理标题
                if re.match(r'chapter_[\d_]+$', attr):
                    if re.match(r'chapter_\d{2}$', attr):
                        # 一级标题
                        if attr not in chapter_list:
                            chapter_list.append(attr)
                    else:
                        # 非一级标题
                        if not subchapter_dict.get(attr):
                            subchapter_dict[attr] = set()
                        if not paragraph_dict.get(attr):
                            paragraph_dict[attr] = list()
                        self.mount_parent_chapter(subchapter_dict, attr)
                    continue

                # 处理段落
                for regex_rule in regex_rules:
                    result = re.match(regex_rule, attr)
                    if result:
                        chapter = result.group(1)
                        self.mount_parent_chapter(subchapter_dict, chapter)
                        if not paragraph_dict.get(chapter):
                            paragraph_dict[chapter] = list()
                        paragraph_dict[chapter].append(attr)

        # 排序
        chapter_list = sorted(chapter_list)
        for k, v in subchapter_dict.items():
            subchapter_dict[k] = sorted(v)
        self.chapter_list = chapter_list
        self.subchapter_dict = subchapter_dict
        self.paragraph_dict = paragraph_dict

    def _handle_chapter(self):
        """处理章节"""
        for chapter in self.chapter_list:
            handler = getattr(self, chapter)
            try:
                handler()
            except Exception as e:
                raise ValueError('处理章节标题chapter[{}]处理失败: {}'.format(chapter, e))
            self._handle_subchapters(chapter)

    def mount_parent_chapter(self, subchapter_dict, chapter):
        """挂载父类章节"""
        res = chapter.split('_')
        if len(res) > 2:
            parent_chapter = '_'.join(res[:-1])
            if not subchapter_dict.get(parent_chapter):
                subchapter_dict[parent_chapter] = set()
            subchapter_dict[parent_chapter].add(chapter)
            self.mount_parent_chapter(subchapter_dict, parent_chapter)

    def get_paragraph_data(self, function_name: str = None):
        """获取章段落对应数据"""
        if not function_name:
            function_name = sys._getframe(1).f_code.co_name
        data = self.data.get(function_name)
        return data if data else dict()

    def _handle_subchapters(self, chapter: str):
        """处理子章节"""
        subchapters = self.subchapter_dict.get(chapter, list())
        for subchapter in subchapters:
            handler = getattr(self, subchapter)
            try:
                handler()
            except Exception as e:
                raise ValueError('处理章节标题chapter[{}]处理失败: {}'.format(subchapter, e))

    def _handle_paragraphs(self, chapter):
        """段落"""
        paragraphs = self.paragraph_dict.get(chapter, list())
        for paragraph in paragraphs:
            handler = getattr(self, paragraph)
            try:
                handler()
            except Exception as e:
                raise ValueError('处理段落paragraph[{}]处理失败: {}'.format(paragraph, e))

    def _clear(self):
        """清理临时文件"""
        for f in self._clear_files:
            try:
                os.remove(f)
            except Exception as e:
                pass

    def _handle_filename(self):
        """处理文件名"""
        current = datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')
        if self.filename.endswith('.docx'):
            self.filename = '{}_{}.docx'.format(self.filename[:-5], current)
        elif self.filename.endswith('.doc'):
            self.filename = '{}_{}.doc'.format(self.filename[:-5], current)
        else:
            self.filename += '_.docx'
