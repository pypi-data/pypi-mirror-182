# -*- coding: utf-8 -*-
"""
@File  : picture.py
@IDE   : PyCharm
@Author: Sanbom
@Date  : 2021/7/15
@Desc  : 
"""

from docx.document import Document as DocObject
from docx.shared import Cm

from smartDocx.config import DocStylesBase, PicImgStyleBase
from smartDocx.package.content import Content


class Picture(object):
    """图片"""

    def __init__(self, config: PicImgStyleBase):
        self.config = config
        if not (0 < self.config.zom_rate <= 1):
            self.width = 15
        else:
            self.width = round(15 * self.config.zom_rate, 3)

    def generate(self, doc_obj: DocObject, pic_path, pic_title: str = None,
                 pic_title_style: DocStylesBase = None):
        """
        生成图片
        :param doc_obj:
        :param pic_path: 图片地址
        :param pic_title: 图片标题
        :param pic_title_style: 图片标题样式
        :return:
        """
        if pic_title and pic_title_style:
            content = Content(config=pic_title_style)
            content.generate(doc_obj=doc_obj, text=pic_title)
        paragraph = doc_obj.add_paragraph()
        paragraph.alignment = self.config.alignment.value
        run = paragraph.add_run()
        run.add_picture(image_path_or_stream=pic_path, width=Cm(self.width))
